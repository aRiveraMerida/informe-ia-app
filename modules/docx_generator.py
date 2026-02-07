# -*- coding: utf-8 -*-
"""
Módulo de generación de documentos DOCX editables.
Convierte markdown a DOCX con formato profesional.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
from datetime import datetime
from typing import Optional, Dict, Any, List
import re
import logging

logger = logging.getLogger(__name__)


class DOCXReportGenerator:
    """Genera documentos DOCX profesionales editables."""
    
    def __init__(
        self,
        client_name: str,
        period: str,
        report_title: str = "Informe Ejecutivo"
    ):
        self.client_name = client_name
        self.period = period
        self.report_title = report_title
        self.document = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Configura estilos personalizados del documento."""
        styles = self.document.styles
        
        # Título principal — Movimer
        title_style = styles['Title']
        title_style.font.size = Pt(28)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(112, 174, 0)
        
        # Heading 1
        h1_style = styles['Heading 1']
        h1_style.font.size = Pt(18)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(112, 174, 0)
        
        # Heading 2
        h2_style = styles['Heading 2']
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(61, 107, 0)
        
        # Heading 3
        h3_style = styles['Heading 3']
        h3_style.font.size = Pt(12)
        h3_style.font.bold = True
        h3_style.font.color.rgb = RGBColor(74, 140, 0)
    
    def generate(
        self,
        analysis_text: str,
        metadata: Dict[str, Any],
        chart_images: Optional[List[tuple]] = None,
        quantitative_analysis: Optional[str] = None,
        report_sections: Optional[List] = None,
    ) -> bytes:
        """
        Genera un DOCX unificado profesional.

        Si se proporcionan report_sections (List[ReportSection]), se usa el
        flujo de secciones con gráficos inline. Si no, se usa el flujo legacy.

        Estructura:
          1. Portada
          2. Secciones del informe con gráficos inline
          3. Anexo estadístico
          4. Nota de generación
        """
        # 1. Portada
        self._add_cover_page(metadata)

        # 2. Informe con gráficos inline
        if report_sections:
            self._add_sections_with_charts(report_sections)
        else:
            # Flujo legacy
            self._parse_markdown(analysis_text)
            if chart_images:
                self._add_charts_section(chart_images)

        # 3. Nota de generación
        self._add_footer_info(metadata)

        buffer = BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _add_sections_with_charts(self, report_sections: List):
        """Agrega secciones del informe con gráficos inline."""
        for section in report_sections:
            # Heading
            if section.heading:
                self.document.add_heading(section.heading, level=2)

            # Texto de la sección
            if section.text:
                self._parse_markdown(section.text)

            # Gráficos inline
            if section.chart_images:
                self.document.add_paragraph()  # Espacio
                for chart_title, png_bytes in section.chart_images:
                    try:
                        img_stream = BytesIO(png_bytes)
                        self.document.add_picture(img_stream, width=Inches(5.8))
                        # Pie de gráfico
                        caption = self.document.add_paragraph()
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = caption.add_run(chart_title)
                        run.font.size = Pt(8)
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(128, 128, 128)
                    except Exception as e:
                        logger.warning("No se pudo insertar gráfico '%s': %s", chart_title, e)
                self.document.add_paragraph()  # Espacio tras gráficos
    
    def _add_cover_page(self, metadata: Dict[str, Any]):
        """Agrega portada al documento."""
        # Espacio superior
        for _ in range(3):
            self.document.add_paragraph()
        
        # Título
        title = self.document.add_heading(self.report_title.upper(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Línea divisoria
        self.document.add_paragraph("─" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtítulo/Cliente
        subtitle = self.document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(self.client_name)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(61, 107, 0)
        
        # Espacio
        self.document.add_paragraph()
        
        # Metadata
        meta_para = self.document.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.add_run(f"Periodo: {self.period}\n").font.size = Pt(12)
        
        if 'total_records' in metadata:
            meta_para.add_run(f"Total registros: {metadata['total_records']:,}\n").font.size = Pt(12)
        
        fecha = datetime.now().strftime("%d de %B de %Y")
        meta_para.add_run(f"Fecha: {fecha}").font.size = Pt(12)
        
        # Salto de página
        self.document.add_page_break()
    
    def _parse_markdown(self, text: str):
        """Convierte markdown a elementos DOCX."""
        lines = text.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # Heading 1 (#)
            if line.startswith('# ') and not line.startswith('## '):
                heading_text = line[2:].strip()
                self.document.add_heading(heading_text, level=1)
            
            # Heading 2 (##)
            elif line.startswith('## '):
                heading_text = line[3:].strip()
                self.document.add_heading(heading_text, level=2)
            
            # Heading 3 (###)
            elif line.startswith('### '):
                heading_text = line[4:].strip()
                self.document.add_heading(heading_text, level=3)
            
            # Tabla markdown
            elif '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
                table_lines = [line]
                i += 1
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                self._add_table(table_lines)
                continue  # Ya avanzamos i
            
            # Bullet points
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                bullet_text = line.strip()[2:]
                bullet_text = self._clean_markdown(bullet_text)
                para = self.document.add_paragraph(bullet_text, style='List Bullet')
            
            # Línea horizontal
            elif line.strip() in ['---', '___', '***']:
                para = self.document.add_paragraph('─' * 60)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Texto con negrita (ej: **texto**)
            elif line.strip().startswith('**') and line.strip().endswith('**'):
                text_content = line.strip()[2:-2]
                para = self.document.add_paragraph()
                run = para.add_run(text_content)
                run.bold = True
            
            # Texto normal
            elif line.strip():
                clean_text = self._clean_markdown(line)
                para = self.document.add_paragraph()
                self._add_formatted_text(para, clean_text)
            
            # Línea vacía
            else:
                self.document.add_paragraph()
            
            i += 1
    
    def _clean_markdown(self, text: str) -> str:
        """Limpia formato markdown básico para texto plano."""
        # Eliminar ** pero mantener el texto
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'__(.+?)__', r'\1', text)
        # Eliminar * e _ para cursiva
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'_(.+?)_', r'\1', text)
        return text
    
    def _add_formatted_text(self, paragraph, text: str):
        """Agrega texto con formato (negrita, cursiva) al párrafo."""
        # Buscar patrones de negrita **texto**
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)
    
    def _add_table(self, table_lines: List[str]):
        """Convierte tabla markdown a tabla DOCX."""
        # Parsear líneas de tabla
        rows = []
        for line in table_lines:
            if '---' in line:
                continue  # Saltar separador
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                rows.append(cells)
        
        if not rows:
            return
        
        # Crear tabla
        num_cols = len(rows[0])
        table = self.document.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Llenar tabla
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = self._clean_markdown(cell_text)
                    
                    # Formato header (primera fila)
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                        # Fondo azul para header
                        from docx.oxml.ns import nsdecls
                        from docx.oxml import parse_xml
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="70AE00"/>')
                        cell._tc.get_or_add_tcPr().append(shading)
        
        # Espacio después de tabla
        self.document.add_paragraph()
    
    def _add_charts_section(self, chart_images: List[tuple]):
        """Agrega visualizaciones integradas tras el informe ejecutivo."""
        self.document.add_page_break()
        self.document.add_heading("Visualizaciones de Datos", level=1)
        intro = self.document.add_paragraph()
        run = intro.add_run(
            "Los siguientes gráficos complementan el análisis presentado "
            "en las secciones anteriores."
        )
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)
        self.document.add_paragraph()

        for title, png_bytes in chart_images:
            self.document.add_heading(title, level=3)
            try:
                img_stream = BytesIO(png_bytes)
                self.document.add_picture(img_stream, width=Inches(6))
                self.document.add_paragraph()
            except Exception as e:
                logger.warning(f"No se pudo insertar gráfico '{title}': {e}")

    def _add_footer_info(self, metadata: Dict[str, Any]):
        """Agrega información de generación al final."""
        self.document.add_paragraph()
        self.document.add_paragraph('─' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer = self.document.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer_text = f"Informe generado automáticamente con IA \n"
        footer_text += f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        
        run = footer.add_run(footer_text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
